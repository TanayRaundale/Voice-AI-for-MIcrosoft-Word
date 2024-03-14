import docx
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
import pyttsx3

def speak(audio):
    engine = pyttsx3.init('sapi5')
    voices = engine.getProperty('voices')
    engine.setProperty('voices', voices[0].id)
    engine.say(audio)
    engine.runAndWait()

def create_empty_table(num_rows, num_cols, docname):
    document = docx.Document()
    # Create a table with the specified number of rows and columns
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Set the cell borders
    for row in table.rows:
        for cell in row.cells:
            cell.text = " "
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
            for border in cell._element.xpath('.//w:tcBorders'):
                border.attrib['w:val'] = 'single'
                border.attrib['w:sz'] = '4'  # Border size (optional)

    document.add_paragraph(" ")
    document.save(docname)
    speak("Table added successfully")

if __name__ == "__main__":
    create_empty_table(4,4,"bappa.docx")