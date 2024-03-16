import requests
import os
import pyttsx3
from docx import Document
from docx.shared import Inches

def speak(audio):
    engine = pyttsx3.init('sapi5')
    voices = engine.getProperty('voices')
    engine.setProperty('voices', voices[0].id)
    engine.say(audio)
    engine.runAndWait()

def generate_images(topic):
    PIXABAY_API_KEY = '41925790-ffb825eb27f460e8ccf41f878'
    url = f'https://pixabay.com/api/?key={PIXABAY_API_KEY}&q={topic}&image_type=photo'
    response = requests.get(url)
    data = response.json()
    urls = []

    if 'hits' in data:
        for hit in data['hits']:
            urls.append(hit['largeImageURL'])
    else:
        print('Error in API response:', data['message'])

    # Create directory to store images
    os.makedirs(topic, exist_ok=True)

    # Create a Word document
    doc = Document()
    doc.add_heading(f'Images for "{topic}"', level=1)

    for index, img_link in enumerate(urls):
        if index >= 10:  # Limit to first 5 images
            break

        # Download image
        img_data = requests.get(img_link).content

        # Save image to file
        img_path = os.path.join(topic, f"image_{index + 1}.jpg")
        with open(img_path, 'wb') as f:
            f.write(img_data)

        # Insert image into Word document
        doc.add_picture(img_path, width=Inches(5))

    # Save the Word document
    doc_file = f"{topic}_images.docx"
    doc.save(doc_file)

    speak(f"Sir, I have generated the Word document with images on the topic '{topic}'.")
    speak(f"I will open the '{doc_file}' file for you.")
    os.startfile(doc_file)


