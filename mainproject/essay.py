from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from selenium.webdriver.common.by import By
import docx
import re
from docx.shared import RGBColor

import os

def generate_essay(search_string):

    doc = docx.Document()


    options = webdriver.FirefoxOptions()
    options.add_argument("-headless")
    driver = webdriver.Firefox(options=options)
    driver.maximize_window()
    f_string = search_string.replace(" ", "_")
    driver.get(f"https://www.wikipedia.org/wiki/{f_string}")



    if "Wikipedia does not have an article with this exact name." in driver.page_source:
        print("sorry we could not find a wikipedia for your search:(")
    else:
        body = driver.find_element(By.ID,"bodyContent")

        p_tags = body.find_elements(By.TAG_NAME,"p")
        head = doc.add_heading(f"Essay on {search_string}",2)

        run = head.runs[0]
        font = run.font
        font.color.rgb = RGBColor(92, 2, 35)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for ptag in p_tags:

           # print(ptag.text)
            pattern = re.compile(r'\[\d+\]')


           # Use the sub() function to replace all occurrences of the pattern with an empty string
            cleaned_text = re.sub(pattern, '', ptag.text)
            para = doc.add_paragraph(cleaned_text)
            para.alignment = 3
        path = f"C:\\Users\\user\\Desktop\\fyp\\{search_string}.docx"
        doc.save(f"{search_string}.docx")

        path = f"C:\\Users\\user\\PycharmProjects\\pythonProject\\{search_string}.docx"
        os.startfile(path)



if __name__== "__main__":
    generate_essay("Albert Einstein")