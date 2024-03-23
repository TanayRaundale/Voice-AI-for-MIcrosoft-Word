from docx import Document
from docx.shared import Pt
import pyttsx3
import os

def speak(audio):
    engine = pyttsx3.init('sapi5')
    voices = engine.getProperty('voices')
    engine.setProperty('voices', voices[0].id)
    engine.setProperty('rate', 160)
    engine.say(audio)
    engine.runAndWait()


def set_margins(doc, top, right, bottom, left):
    sections = doc.sections

    for section in sections:
        section.top_margin = Pt(top)
        section.right_margin = Pt(right)
        section.bottom_margin = Pt(bottom)
        section.left_margin = Pt(left)
    speak("Margin Added Successfully")


def add_margins(docname):

    doc = Document(docname)   # Create a new Word document

    # Set margins (in points)
    top_margin = 1  # 1/2 inch
    right_margin = 1.2
    bottom_margin = 1.75                
    left_margin = 1.00

    set_margins(doc, top_margin, right_margin, bottom_margin, left_margin)

    # doc.add_heading('Document generated with Default Margins CM', level=1)
    # doc.add_paragraph("")
    # doc.add_paragraph('This is an Simple Default Margin Example and generates Document...')
    speak("I will open the modified document for you")
    doc.save(docname)
    os.startfile(f"C:\\Users\\Dell\\Desktop\\final year project\\mainproject\\{docname}") 
